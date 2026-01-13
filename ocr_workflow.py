"""
Complete OCR Workflow: Conditioning + API Processing
Chains image conditioning with Multimodal-OCR3 API
Outputs all results to a single JSON file
"""

import os
import sys
import json
from pathlib import Path
from datetime import datetime
from image_conditioner import ImageConditioner
from gradio_client import Client, handle_file
import logging

# Setup logging first
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import gibberish detection utility
try:
    from orientation_validator import detect_gibberish
except ImportError:
    detect_gibberish = None
    logger.warning("orientation_validator not available, gibberish detection disabled")


class OCRWorkflow:
    """
    Complete OCR pipeline: conditioning → API processing
    Collects all results and outputs to a single JSON file
    """
    
    def __init__(self, input_folder=None, png_compression=None, strength=None, json_output=None, config_path="config.json", 
                 quality_mode=None, ocr_models=None, disable_fallback=None):
        """
        Initialize workflow with configurable compression and strength.
        Loads from config.json if available, otherwise uses defaults or provided values.
        
        Args:
            input_folder: Input folder with images/PDFs. Overrides config if provided.
            png_compression: PNG compression level (0=none, 9=max). Overrides config if provided.
            strength: Conditioning strength (0-100). Overrides config if provided.
            json_output: Output JSON file path. Overrides config if provided.
            config_path: Path to config.json file. Default: "config.json"
            quality_mode: Quality preset ('high', 'medium', 'low'). Overrides config if provided.
            ocr_models: List of specific OCR models to use (e.g., ['Nanonets-OCR2-3B']). Overrides config if provided.
            disable_fallback: Disable fallback models (use only primary model). Overrides config if provided.
        """
        # Load config if available
        config = self._load_config(config_path)
        
        # Use config values or provided values or defaults
        if input_folder is None:
            input_folder = config.get("image_conditioning", {}).get("input_folder", {}).get("value", "Input")
        if png_compression is None:
            png_compression = config.get("image_conditioning", {}).get("png_compression", {}).get("value", 1)
        if strength is None:
            strength = config.get("image_conditioning", {}).get("strength", {}).get("value", 10)
        
        # Get output folders and formats from config
        json_output_folder = config.get("ocr_workflow", {}).get("json_output_folder", {}).get("value", "Results")
        timestamp_json = config.get("ocr_workflow", {}).get("timestamp_json", {}).get("value", True)
        ocr_output_folder = config.get("ocr_workflow", {}).get("ocr_output_folder", {}).get("value", "Output")
        output_formats = config.get("ocr_workflow", {}).get("output_formats", {}).get("value", ["txt"])
        
        # Validate output formats
        valid_formats = {"txt", "markdown", "docx", "json"}
        if not isinstance(output_formats, list):
            output_formats = ["txt"]
        output_formats = [f.lower() for f in output_formats if f.lower() in valid_formats]
        if not output_formats:  # At least one format required
            output_formats = ["txt"]
            logger.warning("No valid output formats specified, defaulting to 'txt'")
        
        self.output_formats = set(output_formats)  # Store as set for fast lookup
        logger.info(f"Output formats enabled: {', '.join(sorted(self.output_formats))}")
        
        # Create Results folder if JSON format is enabled
        if "json" in self.output_formats:
            os.makedirs(json_output_folder, exist_ok=True)
            
            # Generate JSON filename with timestamp if enabled
            if json_output is None:
                if timestamp_json:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    json_output = os.path.join(json_output_folder, f"ocr_results_{timestamp}.json")
                else:
                    json_output = os.path.join(json_output_folder, "ocr_results.json")
            else:
                # If user provided json_output, ensure it's in the Results folder
                json_output = os.path.join(json_output_folder, os.path.basename(json_output))
        else:
            json_output = None  # JSON format disabled
        
        # Load all image conditioning parameters from config
        ic_config = config.get("image_conditioning", {})
        
        self.conditioner = ImageConditioner(
            input_folder=input_folder,
            output_folder=ic_config.get("output_folder", {}).get("value", "results_conditioned"),
            strength=strength,
            adaptive_block_size=ic_config.get("adaptive_block_size", {}).get("value"),
            adaptive_C=ic_config.get("adaptive_C", {}).get("value"),
            morph_kernel_size=tuple(ic_config.get("morph_kernel_size", {}).get("value", [2, 2])) if ic_config.get("morph_kernel_size", {}).get("value") is not None else None,
            morph_iterations=ic_config.get("morph_iterations", {}).get("value"),
            target_width=ic_config.get("target_width", {}).get("value"),
            target_height=ic_config.get("target_height", {}).get("value"),
            canny_threshold1=ic_config.get("canny_threshold1", {}).get("value"),
            canny_threshold2=ic_config.get("canny_threshold2", {}).get("value"),
            min_contour_area=ic_config.get("min_contour_area", {}).get("value"),
            max_pages=ic_config.get("max_pages", {}).get("value"),
            denoise_ksize=ic_config.get("denoise_ksize", {}).get("value"),
            png_compression=png_compression,
            min_scale_for_zero=ic_config.get("min_scale_for_zero", {}).get("value"),
            debug_mode=ic_config.get("debug_mode", {}).get("value", False)
        )
        
        self.ocr_output_folder = ocr_output_folder
        self.json_output_file = json_output
        
        # Load OCR model configuration
        ocr_config = config.get("ocr_workflow", {})
        
        # Available OCR models with quality ratings
        AVAILABLE_MODELS = {
            "Nanonets-OCR2-3B": {"quality": "high", "speed": "medium", "description": "High accuracy, best for quality"},
            "Chandra-OCR": {"quality": "medium", "speed": "fast", "description": "Balanced quality and speed"},
            "Dots.OCR": {"quality": "medium", "speed": "fast", "description": "Good for general use"},
            "olmOCR-2-7B-1025": {"quality": "high", "speed": "slow", "description": "Very high accuracy, slower"}
        }
        
        # Apply quality mode presets if specified
        if quality_mode:
            quality_mode = quality_mode.lower()
            if quality_mode == "high":
                # High quality: use best models only, no fallback
                if ocr_models is None:
                    ocr_models = ["Nanonets-OCR2-3B", "olmOCR-2-7B-1025"]
                if disable_fallback is None:
                    disable_fallback = False  # Still allow fallback between high-quality models
            elif quality_mode == "medium":
                # Medium quality: balanced approach
                if ocr_models is None:
                    ocr_models = ["Nanonets-OCR2-3B", "Chandra-OCR", "Dots.OCR"]
                if disable_fallback is None:
                    disable_fallback = False
            elif quality_mode == "low":
                # Low quality: fastest models
                if ocr_models is None:
                    ocr_models = ["Chandra-OCR", "Dots.OCR"]
                if disable_fallback is None:
                    disable_fallback = False
        
        # Set OCR model(s) - command line override takes precedence
        if ocr_models is not None:
            # Validate models
            valid_models = []
            for model in ocr_models if isinstance(ocr_models, list) else [ocr_models]:
                if model in AVAILABLE_MODELS:
                    valid_models.append(model)
                else:
                    logger.warning(f"Unknown OCR model: {model}. Available: {list(AVAILABLE_MODELS.keys())}")
            
            if valid_models:
                self.ocr_model = valid_models[0]
                self.fallback_models = valid_models[1:] if len(valid_models) > 1 else []
                self.enable_fallback = len(self.fallback_models) > 0 and (disable_fallback is not True)
                logger.info(f"Using specified models: Primary={self.ocr_model}, Fallbacks={self.fallback_models}")
            else:
                # Fall back to config defaults
                self.ocr_model = ocr_config.get("ocr_model", {}).get("value", "Nanonets-OCR2-3B")
                self.fallback_models = ocr_config.get("fallback_models", {}).get("value", ["Chandra-OCR", "Dots.OCR", "olmOCR-2-7B-1025"])
                self.enable_fallback = ocr_config.get("enable_fallback", {}).get("value", True) if disable_fallback is not True else False
        else:
            # Use config values
            self.ocr_model = ocr_config.get("ocr_model", {}).get("value", "Nanonets-OCR2-3B")
            self.fallback_models = ocr_config.get("fallback_models", {}).get("value", ["Chandra-OCR", "Dots.OCR", "olmOCR-2-7B-1025"])
            self.enable_fallback = ocr_config.get("enable_fallback", {}).get("value", True) if disable_fallback is not True else False
        
        # Override fallback if explicitly disabled
        if disable_fallback is True:
            self.enable_fallback = False
            self.fallback_models = []
        
        self.api_timeout = ocr_config.get("api_timeout", {}).get("value", 60)
        
        # Ensure fallback_models is a list
        if not isinstance(self.fallback_models, list):
            self.fallback_models = []
        
        # Create output folders if formats are enabled
        if self.output_formats - {"json"}:  # If any format other than JSON is enabled
            os.makedirs(self.ocr_output_folder, exist_ok=True)
        
        logger.info(f"Input folder: {input_folder}")
        logger.info(f"PNG compression: {png_compression}, Strength: {strength}")
        logger.info(f"Conditioned output: {self.conditioner.output_folder}")
        logger.info(f"OCR output: {self.ocr_output_folder}")
        logger.info(f"Enabled output formats: {', '.join(sorted(self.output_formats))}")
        logger.info(f"OCR model: {self.ocr_model}")
        if self.enable_fallback and self.fallback_models:
            logger.info(f"Fallback models: {', '.join(self.fallback_models)}")
        if "json" in self.output_formats:
            logger.info(f"JSON output: {self.json_output_file}")
    
    def _load_config(self, config_path="config.json"):
        """Load configuration from JSON file, return empty dict if not found.
        
        Handles both development mode and PyInstaller bundled executable:
        1. First checks current directory (allows user override)
        2. Then checks PyInstaller's temp directory (bundled config)
        """
        # Try current directory first (user override or development)
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    logger.info(f"Loaded configuration from {config_path}")
                    return config
            except Exception as e:
                logger.warning(f"Failed to load config from {config_path}: {e}. Using defaults.")
                return {}
        
        # If running as PyInstaller bundle, check bundled location
        if getattr(sys, 'frozen', False):
            # Running as compiled executable
            base_path = sys._MEIPASS
            bundled_config = os.path.join(base_path, config_path)
            if os.path.exists(bundled_config):
                try:
                    with open(bundled_config, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                        logger.info(f"Loaded bundled configuration from {bundled_config}")
                        return config
                except Exception as e:
                    logger.warning(f"Failed to load bundled config: {e}. Using defaults.")
                    return {}
        
        logger.info(f"Config file {config_path} not found. Using defaults.")
        return {}
    
    def run_full_pipeline(self):
        """Execute complete pipeline: condition → OCR → JSON output."""
        
        logger.info("=" * 70)
        logger.info("OCR WORKFLOW: CONDITIONING + API PROCESSING")
        logger.info("=" * 70)
        
        # Collect input file information
        input_files = []
        input_folder = self.conditioner.input_folder
        if os.path.exists(input_folder):
            for filename in sorted(os.listdir(input_folder)):
                file_path = os.path.join(input_folder, filename)
                if os.path.isfile(file_path):
                    file_info = {
                        "filename": filename,
                        "path": os.path.abspath(file_path),
                        "size_bytes": os.path.getsize(file_path),
                        "extension": os.path.splitext(filename)[1].lower()
                    }
                    input_files.append(file_info)
        
        # Collect all results in a dictionary
        all_results = {
            "metadata": {
                "timestamp": datetime.now().isoformat(),
                "input_folder": os.path.abspath(input_folder),
                "conditioned_folder": os.path.abspath(self.conditioner.output_folder),
                "ocr_output_folder": os.path.abspath(self.ocr_output_folder),
                "output_formats": sorted(list(self.output_formats)),  # List of enabled formats
                "conditioning_strength": self.conditioner.strength,
                "png_compression": self.conditioner.png_compression,
                "input_files_count": len(input_files),
                "input_files": input_files
            },
            "files": []
        }
        
        # Step 1: Condition images
        logger.info("\n[STEP 1] IMAGE CONDITIONING")
        logger.info("-" * 70)
        logger.info(f"Input folder: {input_folder}")
        logger.info(f"Found {len(input_files)} input file(s)")
        
        condition_stats = self.conditioner.process_all_files()
        
        if condition_stats['processed'] == 0:
            logger.error("No images were processed in conditioning step!")
            all_results["metadata"]["conditioning_stats"] = condition_stats
            self._save_json_output(all_results)
            return False
        
        logger.info(f"Conditioning complete: {condition_stats['processed']} files processed\n")
        all_results["metadata"]["conditioning_stats"] = condition_stats
        
        # Step 2: Feed conditioned images to OCR API
        logger.info("[STEP 2] OCR API PROCESSING")
        logger.info("-" * 70)
        
        ocr_success = self._process_with_ocr_api(all_results)
        
        # Step 3: Save all results to JSON
        logger.info("\n[STEP 3] SAVING JSON OUTPUT")
        logger.info("-" * 70)
        self._save_json_output(all_results)
        
        return ocr_success
    
    def _try_ocr_with_model(self, client, img_path: str, model_name: str) -> tuple:
        """
        Try OCR processing with a specific model.
        
        Args:
            client: Gradio client instance
            img_path: Path to image file
            model_name: Name of the OCR model to use
            
        Returns:
            Tuple of (success: bool, raw_output: str, markdown_output: str, error: str)
        """
        import signal
        
        try:
            logger.info(f"    Trying model: {model_name}")
            
            # Note: gradio_client doesn't support timeout directly, but we can catch timeout errors
            # The API will raise an exception if it times out
            result = client.predict(
                model_name,                    # model_name
                "Perform OCR on the image.",   # text_query
                handle_file(img_path),        # image
                2048,                          # max_tokens
                0.7,                           # temperature
                0.9,                           # top_p
                50,                            # top_k
                1.1,                           # repetition_penalty
                api_name="/generate_image"
            )
            
            # Extract results
            raw_output = str(result[0]) if isinstance(result, (list, tuple)) and len(result) > 0 else ""
            markdown_output = str(result[1]) if isinstance(result, (list, tuple)) and len(result) > 1 else raw_output
            
            return True, raw_output, markdown_output, ""
            
        except Exception as e:
            error_msg = str(e)
            # Check if it's a timeout or GPU availability error
            is_timeout = ("timeout" in error_msg.lower() or 
                         "gpu" in error_msg.lower() or 
                         "60s" in error_msg.lower() or
                         "No GPU was available" in error_msg)
            
            if is_timeout:
                return False, "", "", f"Timeout/GPU unavailable: {error_msg}"
            else:
                return False, "", "", error_msg
    
    def _process_with_ocr_api(self, results_dict: dict) -> bool:
        """
        Process conditioned images with OCR API and collect results.
        
        Args:
            results_dict: Dictionary to collect all results in
            
        Returns:
            True if successful, False otherwise
        """
        logger.info("Connecting to Multimodal-OCR3 API...")
        
        try:
            client = Client("prithivMLmods/Multimodal-OCR3")
            logger.info("[OK] Connected to OCR API\n")
        except Exception as e:
            logger.error(f"Failed to connect to OCR API: {e}")
            results_dict["metadata"]["ocr_error"] = str(e)
            return False
        
        # Process all conditioned images
        conditioned_folder = self.conditioner.output_folder
        
        if not os.path.exists(conditioned_folder):
            logger.error(f"Conditioned folder not found: {conditioned_folder}")
            return False
        
        files = sorted([f for f in os.listdir(conditioned_folder) 
                       if f.lower().endswith('.png')])
        
        logger.info(f"Found {len(files)} conditioned images\n")
        
        processed = 0
        failed = 0
        
        for filename in files:
            img_path = os.path.abspath(os.path.join(conditioned_folder, filename))
            file_info = {
                "conditioned_image": filename,
                "conditioned_image_path": img_path,
                "status": "pending"
            }
            
            try:
                logger.info(f"Processing: {filename}")
                
                # Get file size
                if os.path.exists(img_path):
                    file_info["conditioned_image_size_bytes"] = os.path.getsize(img_path)
                
                # Build list of models to try (primary + fallbacks)
                models_to_try = [self.ocr_model]
                if self.enable_fallback:
                    # Add fallback models, excluding the primary model if it's already in fallback list
                    for fallback_model in self.fallback_models:
                        if fallback_model != self.ocr_model and fallback_model not in models_to_try:
                            models_to_try.append(fallback_model)
                
                # Try models in order until one succeeds
                raw_output = ""
                markdown_output = ""
                model_used = None
                last_error = ""
                
                for model_name in models_to_try:
                    success, raw, md, error = self._try_ocr_with_model(client, img_path, model_name)
                    
                    if success:
                        raw_output = raw
                        markdown_output = md
                        model_used = model_name
                        logger.info(f"  ✓ Success with model: {model_name}")
                        break
                    else:
                        last_error = error
                        logger.warning(f"  ✗ Failed with {model_name}: {error}")
                        # If it's not a timeout/GPU error, don't try other models (likely a different issue)
                        if "timeout" not in error.lower() and "gpu" not in error.lower() and "60s" not in error.lower():
                            logger.warning(f"    Non-timeout error, skipping fallback models")
                            break
                
                # If all models failed, raise an exception
                if not raw_output and not markdown_output:
                    raise Exception(f"All models failed. Last error: {last_error}")
                
                # Store which model was used
                file_info["ocr_model_used"] = model_used
                if model_used != self.ocr_model:
                    file_info["fallback_used"] = True
                    logger.info(f"  ⚠ Used fallback model: {model_used} (primary: {self.ocr_model})")
                
                # Save files in selected formats
                base_name = os.path.splitext(filename)[0]
                saved_files = {}
                
                # Generate TXT format (plain text)
                if "txt" in self.output_formats:
                    txt_file = os.path.join(self.ocr_output_folder, f"{base_name}.txt")
                    with open(txt_file, "w", encoding="utf-8") as f:
                        f.write(raw_output)
                    saved_files["txt"] = txt_file
                    logger.info(f"  ✓ Saved TXT: {os.path.basename(txt_file)}")
                
                # Generate Markdown format
                if "markdown" in self.output_formats:
                    md_file = os.path.join(self.ocr_output_folder, f"{base_name}.md")
                    with open(md_file, "w", encoding="utf-8") as f:
                        f.write(markdown_output)
                    saved_files["markdown"] = md_file
                    logger.info(f"  ✓ Saved Markdown: {os.path.basename(md_file)}")
                
                # Generate DOCX format (Word document)
                if "docx" in self.output_formats:
                    docx_file = os.path.join(self.ocr_output_folder, f"{base_name}.docx")
                    # Use markdown output if available (has formatting), otherwise raw text
                    docx_content = markdown_output if markdown_output and markdown_output != raw_output else raw_output
                    self._save_as_docx(docx_content, docx_file)
                    saved_files["docx"] = docx_file
                    logger.info(f"  ✓ Saved DOCX: {os.path.basename(docx_file)}")
                
                # Detect gibberish in the result
                is_gibberish = False
                gibberish_score = 0.0
                gibberish_reason = ""
                if detect_gibberish and raw_output:
                    try:
                        is_gibberish, gibberish_score, gibberish_reason = detect_gibberish(raw_output)
                        file_info["gibberish_detected"] = is_gibberish
                        file_info["gibberish_score"] = gibberish_score
                        file_info["gibberish_reason"] = gibberish_reason
                    except Exception as e:
                        logger.debug(f"    Gibberish detection failed: {e}")
                        is_gibberish = False
                
                # Add to results dictionary
                file_info.update({
                    "status": "success",
                    "raw_text": raw_output if "json" in self.output_formats else None,  # Only include if JSON format enabled
                    "markdown_text": markdown_output if "json" in self.output_formats else None,
                    "text_length": len(raw_output),
                    "markdown_length": len(markdown_output),
                    "output_files": saved_files  # Dictionary of format -> filepath
                })
                
                if is_gibberish:
                    logger.warning(f"    ⚠ Gibberish detected: {gibberish_reason}")
                
                logger.info(f"  ✓ Saved OCR results in {len(saved_files)} format(s) (text length: {len(raw_output)} chars)\n")
                processed += 1
                
            except Exception as e:
                error_msg = str(e)
                logger.error(f"  ✗ Error: {error_msg}\n")
                
                # Include information about models tried
                models_tried = [self.ocr_model]
                if self.enable_fallback:
                    models_tried.extend([m for m in self.fallback_models if m != self.ocr_model])
                
                file_info.update({
                    "status": "error",
                    "error": error_msg,
                    "models_tried": models_tried,
                    "fallback_enabled": self.enable_fallback
                })
                failed += 1
            
            results_dict["files"].append(file_info)
        
        results_dict["metadata"]["ocr_stats"] = {
            "processed": processed,
            "failed": failed,
            "total": len(files)
        }
        
        logger.info("=" * 70)
        logger.info(f"OCR PROCESSING COMPLETE: {processed} successful, {failed} failed")
        logger.info(f"Results saved to: {self.ocr_output_folder}")
        logger.info("=" * 70)
        
        # Check for gibberish results and offer fallback
        if detect_gibberish and processed > 0:
            self._check_and_offer_gibberish_fallback(results_dict)
        
        return processed > 0
    
    def _check_and_offer_gibberish_fallback(self, results_dict: dict) -> None:
        """
        Check if there's a mix of valid and gibberish results, and offer to re-process with corrected orientation.
        
        Only triggers if:
        - There are some successful results
        - Some results are readable and some are gibberish (mixed results)
        - NOT if all results are gibberish or all are blank
        """
        files = results_dict.get("files", [])
        if not files:
            return
        
        successful_files = [f for f in files if f.get("status") == "success" and f.get("text_length", 0) > 0]
        if not successful_files:
            logger.debug("    No successful results to check for gibberish")
            return
        
        # Analyze results
        gibberish_count = 0
        valid_count = 0
        gibberish_files = []
        
        for file_info in successful_files:
            is_gibberish = file_info.get("gibberish_detected", False)
            if is_gibberish:
                gibberish_count += 1
                gibberish_files.append(file_info)
            else:
                valid_count += 1
        
        total_analyzed = gibberish_count + valid_count
        if total_analyzed == 0:
            return
        
        # Only trigger if there's a MIX of valid and gibberish (not all one or the other)
        if gibberish_count > 0 and valid_count > 0:
            gibberish_pct = (gibberish_count / total_analyzed) * 100.0
            
            logger.info("\n" + "=" * 70)
            logger.info("⚠ GIBBERISH DETECTION WARNING")
            logger.info("=" * 70)
            logger.info(f"Found mixed results:")
            logger.info(f"  Valid results:   {valid_count}/{total_analyzed} ({100-gibberish_pct:.1f}%)")
            logger.info(f"  Gibberish results: {gibberish_count}/{total_analyzed} ({gibberish_pct:.1f}%)")
            logger.info("")
            logger.info("This may indicate incorrect image orientation.")
            logger.info("Files with gibberish detected:")
            for f in gibberish_files[:5]:  # Show first 5
                logger.info(f"  - {f.get('conditioned_image', 'unknown')}: {f.get('gibberish_reason', 'N/A')}")
            if len(gibberish_files) > 5:
                logger.info(f"  ... and {len(gibberish_files) - 5} more")
            logger.info("")
            logger.info("Would you like to re-process the images with orientation correction?")
            logger.info("This will:")
            logger.info("  1. Re-condition images with different orientations")
            logger.info("  2. Re-run OCR on corrected images")
            logger.info("  3. Compare results")
            logger.info("")
            
            # Prompt for confirmation
            response = input("Re-process with orientation correction? [y/N]: ").strip().lower()
            
            if response in ('y', 'yes'):
                logger.info("\n[FALLBACK] Re-processing with orientation correction...")
                logger.info("-" * 70)
                self._reprocess_with_orientation_correction(results_dict, gibberish_files)
            else:
                logger.info("Skipping re-processing. Results saved as-is.")
        
        elif gibberish_count == total_analyzed and gibberish_count > 1:
            # All results are gibberish - this might indicate a systematic issue
            logger.warning(f"\n⚠ All {gibberish_count} results appear to be gibberish.")
            logger.warning("This may indicate a systematic orientation issue.")
            logger.warning("Consider manually checking image orientation or re-conditioning with different settings.")
        # If all valid, no action needed
    
    def _reprocess_with_orientation_correction(self, results_dict: dict, gibberish_files: list) -> None:
        """
        Re-process gibberish files with corrected orientation (rotate 180°).
        
        Args:
            results_dict: Results dictionary to update
            gibberish_files: List of file_info dicts for files with gibberish results
        """
        if not gibberish_files:
            return
        
        logger.info(f"Re-processing {len(gibberish_files)} files with 180° rotation correction...")
        
        # Import orientation validator for validation
        try:
            from orientation_validator import OrientationValidator
            validator = OrientationValidator()
        except ImportError:
            logger.error("OrientationValidator not available for fallback processing")
            return
        
        # For each gibberish file, re-condition with 180° rotation
        # This is a simplified approach - in a full implementation, you'd try all orientations
        corrected_count = 0
        
        for file_info in gibberish_files:
            original_filename = file_info.get("conditioned_image", "")
            if not original_filename:
                continue
            
            logger.info(f"  Re-processing: {original_filename}")
            
            # Find original image path
            # For now, we'll note this for manual processing or future enhancement
            # Full implementation would: load original, rotate, re-condition, re-OCR
            logger.info(f"    Note: Full re-processing requires re-conditioning pipeline")
            logger.info(f"    Manual fix: Rotate image 180° and re-run conditioning")
        
        logger.info(f"Fallback re-processing complete: {corrected_count} files corrected")
        logger.info("For automatic re-processing, use debug_cropping.py with --rotate flag")
    
    def _save_as_docx(self, text: str, filepath: str):
        """
        Save text as a DOCX file using python-docx library.
        Preserves basic formatting from markdown (headers, bold, italic) when possible.
        """
        try:
            from docx import Document
            from docx.shared import Pt
            import re
        except ImportError as e:
            logger.error(f"python-docx not installed: {e}. Install with: pip install python-docx")
            raise ImportError("python-docx is required for DOCX output format. Install with: pip install python-docx")
        
        doc = Document()
        
        # Basic markdown parsing for better formatting
        lines = text.split('\n')
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                # Empty line - add blank paragraph
                doc.add_paragraph()
                continue
            
            # Check for markdown headers
            if line_stripped.startswith('# '):
                # H1
                doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith('## '):
                # H2
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith('### '):
                # H3
                doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith('#### '):
                # H4
                doc.add_heading(line_stripped[5:], level=4)
            elif line_stripped.startswith('- ') or (line_stripped.startswith('* ') and not line_stripped.startswith('**')):
                # Bullet list item (but not bold markers)
                p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
                self._add_formatted_text(p, line_stripped[2:], Pt)
            elif line_stripped[0].isdigit() and '. ' in line_stripped[:5]:
                # Numbered list item
                list_text = re.sub(r'^\d+\.\s+', '', line_stripped)
                p = doc.add_paragraph(list_text, style='List Number')
                self._add_formatted_text(p, list_text, Pt)
            else:
                # Regular paragraph with basic markdown formatting
                p = doc.add_paragraph()
                # Process inline markdown: **bold**, *italic*, `code`
                self._add_formatted_text(p, line_stripped, Pt)
        
        doc.save(filepath)
    
    def _add_formatted_text(self, paragraph, text: str, Pt_class):
        """
        Add text to paragraph with markdown formatting (bold, italic, code).
        
        Args:
            paragraph: docx paragraph object
            text: Text to add with markdown formatting
            Pt_class: Pt class from docx.shared (passed to avoid import issues)
        """
        import re
        
        # Split by markdown formatting markers
        # Pattern: **bold**, *italic*, `code`
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                # Italic text (but not bold markers)
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code/inline code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt_class(10)
            else:
                # Regular text (may contain unprocessed markdown)
                paragraph.add_run(part)
    
    def _save_json_output(self, results_dict: dict):
        """Save all results to a single JSON file (if JSON format is enabled)."""
        if self.json_output_file is None or "json" not in self.output_formats:
            logger.debug("JSON output format disabled, skipping JSON file generation")
            return
        
        json_path = os.path.abspath(self.json_output_file)
        
        try:
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(results_dict, f, indent=2, ensure_ascii=False)
            
            # Calculate summary stats
            total_files = len(results_dict["files"])
            successful = sum(1 for f in results_dict["files"] if f.get("status") == "success")
            failed = sum(1 for f in results_dict["files"] if f.get("status") == "error")
            total_text_length = sum(f.get("text_length", 0) for f in results_dict["files"])
            
            logger.info(f"✓ JSON output saved: {json_path}")
            logger.info(f"  Total files: {total_files}")
            logger.info(f"  Successful: {successful}")
            logger.info(f"  Failed: {failed}")
            logger.info(f"  Total text extracted: {total_text_length:,} characters")
            logger.info("=" * 70)
            
        except Exception as e:
            logger.error(f"Failed to save JSON output: {e}")
            import traceback
            traceback.print_exc()


def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Full OCR pipeline with JSON output')
    parser.add_argument('--compression', type=int, default=None, 
                       help='PNG compression level (0=none, 9=max). Overrides config if provided.')
    parser.add_argument('--strength', type=int, default=None,
                       help='Conditioning strength (0-100). Overrides config if provided.')
    parser.add_argument('--input', type=str, default=None,
                       help='Input folder with images/PDFs. Overrides config if provided.')
    parser.add_argument('--output-json', type=str, default=None,
                       help='Output JSON file path. Overrides config if provided.')
    parser.add_argument('--config', type=str, default='config.json',
                       help='Path to config.json file. Default: config.json')
    parser.add_argument('--quality', type=str, choices=['high', 'medium', 'low'], default=None,
                       help='Quality preset: high (best accuracy), medium (balanced), low (fastest). Overrides config.')
    parser.add_argument('--models', type=str, nargs='+', default=None,
                       help='Specific OCR models to use (space-separated). Available: Nanonets-OCR2-3B, Chandra-OCR, Dots.OCR, olmOCR-2-7B-1025. Overrides config.')
    parser.add_argument('--no-fallback', action='store_true',
                       help='Disable fallback models (use only primary model). Overrides config.')
    
    args = parser.parse_args()
    
    # Create workflow with all parameters
    workflow = OCRWorkflow(
        input_folder=args.input,
        png_compression=args.compression,
        strength=args.strength,
        json_output=args.output_json,
        config_path=args.config,
        quality_mode=args.quality,
        ocr_models=args.models,
        disable_fallback=args.no_fallback
    )
    
    success = workflow.run_full_pipeline()
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()

import os
import pytest
from ocr_workflow import OCRWorkflow
from docx import Document

@pytest.fixture
def workflow():
    return OCRWorkflow(input_folder="tests/dataset", handwriting_mode=True)

def test_docx_export_basic(workflow, tmp_path):
    output_file = tmp_path / "test_output.docx"
    test_text = "# Title\nThis is **bold** and *italic*."
    
    workflow._save_as_docx(test_text, str(output_file))
    
    assert output_file.exists()
    doc = Document(str(output_file))
    
    # Check title
    assert len(doc.paragraphs) >= 1
    assert doc.paragraphs[0].text == "Title"
    assert doc.paragraphs[0].style.name == "Heading 1"
    
    # Check formatting
    # The _add_formatted_text splits by markers, so we check runs
    p1 = doc.paragraphs[1]
    texts = [run.text for run in p1.runs]
    assert "This is " in texts
    assert "bold" in texts
    assert " and " in texts
    assert "italic" in texts

def test_docx_export_lists(workflow, tmp_path):
    output_file = tmp_path / "test_lists.docx"
    test_text = "- Item 1\n- Item 2\n1. Number 1\n2. Number 2"
    
    workflow._save_as_docx(test_text, str(output_file))
    
    doc = Document(str(output_file))
    
    # Bullet list
    assert doc.paragraphs[0].style.name == "List Bullet"
    assert doc.paragraphs[0].text == "Item 1"
    
    # Numbered list
    # Note: _save_as_docx uses regex to strip leading digits
    assert doc.paragraphs[2].style.name == "List Number"
    assert doc.paragraphs[2].text == "Number 1"

def test_docx_export_empty(workflow, tmp_path):
    output_file = tmp_path / "empty.docx"
    workflow._save_as_docx("", str(output_file))
    assert output_file.exists()
    doc = Document(str(output_file))
    assert len(doc.paragraphs) == 0

def test_pdf_export_basic(workflow, tmp_path):
    output_file = tmp_path / "test_output.pdf"
    test_text = "# PDF Title\nThis is a PDF test."
    
    workflow._save_as_pdf(test_text, str(output_file))
    
    assert output_file.exists()
    assert output_file.stat().st_size > 500 # Basic check that it's not empty
